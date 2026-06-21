"""
ingest.py

Document ingestion pipeline for the RAG Q&A bot.

Responsibilities:
1. Scan the data/ directory for supported documents (.pdf, .docx).
2. Extract text page-by-page (PDF) or block-by-block (DOCX), tracking source
   metadata for citation purposes.
3. Chunk the extracted text using recursive-style overlapping chunking.
4. Embed each chunk using Google's text-embedding-004 model and persist the
   vectors + metadata to a local on-disk ChromaDB collection.
"""

import os
import sys
import shutil

from pypdf import PdfReader
from docx import Document as DocxDocument
import chromadb
from tqdm import tqdm

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import config
from embeddings import GeminiEmbeddingFunction



# Step 2: Document Ingestion & Text Extraction


def extract_pdf_pages(file_path: str) -> list[dict]:
    """
    Extracts text page-by-page from a PDF, tracking page numbers and file source.
    Returns a list of {"text": ..., "metadata": {"source": ..., "page": ...}}.
    """
    extracted_data = []
    file_name = os.path.basename(file_path)

    try:
        reader = PdfReader(file_path)
        for index, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                
                clean_text = " ".join(text.split())
                extracted_data.append({
                    "text": clean_text,
                    "metadata": {
                        "source": file_name,
                        "page": index + 1  # 1-indexed for human readability
                    }
                })
    except Exception as e:
        print(f"  [ERROR] Failed reading PDF {file_name}: {e}")

    return extracted_data


def extract_docx_pages(file_path: str) -> list[dict]:
    """
    Extracts text from a DOCX file. DOCX has no fixed "pages" in the file
    format itself, so we treat each top-level paragraph/table as part of a
    single logical unit, but we still tag a synthetic page number of 1 so
    citation formatting stays consistent across file types. Headings are
    preserved inline since they carry useful semantic context.
    """
    extracted_data = []
    file_name = os.path.basename(file_path)

    try:
        doc = DocxDocument(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text.strip())

        # Also pull table content (factsheet.docx has a table)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        clean_text = " ".join(full_text.split())

        if clean_text:
            extracted_data.append({
                "text": clean_text,
                "metadata": {
                    "source": file_name,
                    "page": 1  # DOCX has no native pagination; see docstring above
                }
            })
    except Exception as e:
        print(f"  [ERROR] Failed reading DOCX {file_name}: {e}")

    return extracted_data


def scan_and_extract(data_dir: str) -> list[dict]:
    """
    Scans the data directory for supported files and extracts text + metadata
    from each one, dispatching to the correct extractor by file extension.
    """
    all_pages = []

    if not os.path.isdir(data_dir):
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    files = sorted(os.listdir(data_dir))
    supported_files = [f for f in files if f.lower().endswith((".pdf", ".docx"))]

    if not supported_files:
        raise ValueError(
            f"No supported documents (.pdf, .docx) found in {data_dir}. "
            "Add at least one document before running ingestion."
        )

    print(f"Found {len(supported_files)} document(s) to process:")
    for f in supported_files:
        print(f"  - {f}")

    for file_name in supported_files:
        file_path = os.path.join(data_dir, file_name)
        if file_name.lower().endswith(".pdf"):
            pages = extract_pdf_pages(file_path)
        else:  # .docx
            pages = extract_docx_pages(file_path)

        print(f"  Extracted {len(pages)} page(s)/section(s) from {file_name}")
        all_pages.extend(pages)

    return all_pages


# Step 3: Text Chunking Strategy (recursive-style, overlapping)


def chunk_extracted_pages(
    pages: list[dict],
    chunk_size: int = config.CHUNK_SIZE,
    chunk_overlap: int = config.CHUNK_OVERLAP,
) -> list[dict]:
    """
    Splits page-level documents into smaller, overlapping chunks.
    Ensures that source metadata is carried over to every individual chunk.
    """
    chunks = []

    for page in pages:
        text = page["text"]
        metadata = page["metadata"]

        start = 0
        text_length = len(text)

        # Guard against pathological config (overlap >= chunk_size would infinite-loop)
        step = max(chunk_size - chunk_overlap, 1)

        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk_text = text[start:end]

            chunks.append({
                "text": chunk_text,
                "metadata": {
                    "source": metadata["source"],
                    "page": metadata["page"],
                    "chunk_range": f"{start}-{end}"
                }
            })

            start += step

    return chunks



# Step 4: Persisting the Vector Database


def save_to_vector_db(chunks: list[dict], db_path: str = config.DB_DIR, reset: bool = True):
    """
    Embeds text chunks and saves them into a persistent disk-based ChromaDB.

    If reset=True, any existing collection data is wiped first. This keeps
    ingestion idempotent: re-running ingest.py after adding/removing source
    documents won't leave behind stale chunks from deleted files.
    """
    if reset and os.path.isdir(db_path):
        print(f"Resetting existing database at {db_path} ...")
        shutil.rmtree(db_path)
    os.makedirs(db_path, exist_ok=True)

    client = chromadb.PersistentClient(path=db_path)

    embedding_fn = GeminiEmbeddingFunction(
        api_key=config.GEMINI_API_KEY,
        model_name=config.EMBEDDING_MODEL,
        task_type="retrieval_document",
    )

    collection = client.get_or_create_collection(
        name=config.COLLECTION_NAME,
        embedding_function=embedding_fn,
        metadata={"hnsw:space": "cosine"},
    )

    ids = [f"id_{i}" for i in range(len(chunks))]
    documents = [chunk["text"] for chunk in chunks]
    metadatas = [chunk["metadata"] for chunk in chunks]

    # Batch in groups to stay well under embedding API request limits
    batch_size = 50
    for i in tqdm(range(0, len(documents), batch_size), desc="Embedding & storing chunks"):
        collection.add(
            ids=ids[i:i + batch_size],
            documents=documents[i:i + batch_size],
            metadatas=metadatas[i:i + batch_size],
        )

    print(f"Successfully indexed {len(chunks)} chunks in the vector database at {db_path}")



# Entry point


def main():
    config.validate_config()

    print("=" * 60)
    print("STEP 1/3: Extracting text from documents")
    print("=" * 60)
    pages = scan_and_extract(config.DATA_DIR)
    print(f"\nTotal extracted page/section units: {len(pages)}")

    print("\n" + "=" * 60)
    print("STEP 2/3: Chunking extracted text")
    print("=" * 60)
    chunks = chunk_extracted_pages(pages)
    print(f"Total chunks created: {len(chunks)} "
          f"(chunk_size={config.CHUNK_SIZE}, overlap={config.CHUNK_OVERLAP})")

    print("\n" + "=" * 60)
    print("STEP 3/3: Embedding chunks and saving to ChromaDB")
    print("=" * 60)
    save_to_vector_db(chunks, db_path=config.DB_DIR)

    print("\nIngestion complete. Run `streamlit run src/main.py` or `python src/main.py` to query.")


if __name__ == "__main__":
    main()
